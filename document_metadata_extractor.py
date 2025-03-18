import os
import re
import logging
from typing import Dict, Any, Optional, List
import pdfplumber
import docx
from datetime import datetime
import fitz  # PyMuPDF


class DocumentMetadataExtractor:
    """
    Extracts metadata from document files (PDF, DOCX)
    """
    def __init__(self, logging_level: str = 'INFO'):
        """Initialize the metadata extractor."""
        self.logger = self._setup_logger(logging_level)
    
    def extract_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract comprehensive file metadata for a document.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing metadata attributes
        """
        if not os.path.exists(file_path):
            self.logger.error(f"File does not exist: {file_path}")
            return {}
            
        # Get basic file information
        metadata = {
            'file_size': os.path.getsize(file_path),
            'file_size_formatted': self._format_file_size(os.path.getsize(file_path)),
            'creation_date': datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            'modification_date': datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat()
        }
        
        # Local variable, not included in returned metadata
        file_extension = os.path.splitext(file_path)[1].lower()
        
        # Extract format-specific metadata
        try:
            if file_extension == '.pdf':
                pdf_metadata = self._extract_pdf_metadata(file_path)
                metadata.update(pdf_metadata)
            elif file_extension in ['.doc', '.docx']:
                docx_metadata = self._extract_docx_metadata(file_path)
                metadata.update(docx_metadata)
        except Exception as e:
            self.logger.error(f"Error extracting detailed metadata: {str(e)}")
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract PDF-specific metadata using both pdfplumber and PyMuPDF for comprehensive results.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Dictionary with PDF metadata
        """
        pdf_metadata = {}
        
        # First use pdfplumber for certain metadata
        try:
            with pdfplumber.open(file_path) as pdf:
                # Basic PDF info
                pdf_metadata['page_count'] = len(pdf.pages)
                
                # Extract text for paragraph counting
                all_text = ""
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    all_text += page_text + "\n\n"
                
                # Count paragraphs (separated by double newlines)
                paragraphs = re.split(r'\n\s*\n', all_text)
                pdf_metadata['paragraph_count'] = len([p for p in paragraphs if p.strip()])
                
                # Document info dictionary
                if hasattr(pdf, 'metadata') and pdf.metadata:
                    for key, value in pdf.metadata.items():
                        if key and value and isinstance(value, (str, int, float, bool)):
                            # Use the database column name for common PDF metadata fields
                            if key.lower() == 'author':
                                pdf_metadata['author'] = value
                            elif key.lower() == 'title':
                                pdf_metadata['title'] = value
                            elif key.lower() == 'subject':
                                pdf_metadata['subject'] = value
                            elif key.lower() == 'creator':
                                pdf_metadata['creator'] = value
                            elif key.lower() == 'producer':
                                pdf_metadata['producer'] = value
                            # Skip other fields that don't map to our database columns
        except Exception as e:
            self.logger.warning(f"pdfplumber metadata extraction error: {str(e)}")
        
        # Then use PyMuPDF for more detailed metadata
        try:
            doc = fitz.open(file_path)
            
            # Get document metadata - safely
            try:
                if hasattr(doc, "pdf_version"):
                    pdf_metadata['pdf_version'] = f"PDF {doc.pdf_version}"
            except AttributeError as e:
                self.logger.warning(f"Error accessing pdf_version: {str(e)}")
                
            try:
                if hasattr(doc, "is_encrypted"):
                    # Convert boolean to integer for SQLite compatibility
                    pdf_metadata['is_encrypted'] = 1 if doc.is_encrypted else 0
            except AttributeError as e:
                self.logger.warning(f"Error checking encryption: {str(e)}")
                
            try:
                # Simplistic check for signatures
                has_signatures = False
                if hasattr(doc, "permissions"):
                    has_signatures = bool(doc.permissions & fitz.PDF_PERM_PRINT) == False
                pdf_metadata['has_signatures'] = 1 if has_signatures else 0
            except AttributeError as e:
                self.logger.warning(f"Error checking signatures: {str(e)}")
            
            # Extract font information safely
            try:
                fonts = set()
                
                for page_num, page in enumerate(doc):
                    try:
                        if hasattr(page, "get_fonts"):
                            font_list = page.get_fonts()
                            
                            for font in font_list:
                                if len(font) > 3:  # Make sure the tuple has enough elements
                                    font_name = font[3]  # Font name is at index 3
                                    fonts.add(font_name)
                    except Exception as e:
                        self.logger.warning(f"Error getting fonts from page {page_num}: {str(e)}")
                
                pdf_metadata['fonts_used'] = str(list(fonts))
            except Exception as e:
                self.logger.warning(f"Error extracting fonts: {str(e)}")
            
            # Count images safely
            try:
                image_count = 0
                for page_num, page in enumerate(doc):
                    try:
                        if hasattr(page, "get_images"):
                            image_list = page.get_images()
                            image_count += len(image_list)
                    except Exception as e:
                        self.logger.warning(f"Error counting images on page {page_num}: {str(e)}")
                
                pdf_metadata['image_count'] = image_count
            except Exception as e:
                self.logger.warning(f"Error counting images: {str(e)}")
            
            # Check for form fields safely
            try:
                has_forms = False
                for page in doc:
                    if hasattr(page, "widget_count") and page.widget_count > 0:
                        has_forms = True
                        break
                
                pdf_metadata['has_forms'] = 1 if has_forms else 0
            except Exception as e:
                self.logger.warning(f"Error checking for forms: {str(e)}")
            
            # Check for annotations safely
            try:
                annotation_count = 0
                for page in doc:
                    if hasattr(page, "annots"):
                        annotation_count += len(page.annots())
                
                pdf_metadata['annotation_count'] = annotation_count
            except Exception as e:
                self.logger.warning(f"Error checking annotations: {str(e)}")
            
            # Check if document has TOC/bookmarks safely
            try:
                toc = []
                if hasattr(doc, "get_toc"):
                    toc = doc.get_toc()
                pdf_metadata['has_toc'] = 1 if len(toc) > 0 else 0
                pdf_metadata['toc_items'] = len(toc)
            except Exception as e:
                self.logger.warning(f"Error checking TOC: {str(e)}")
            
            doc.close()
        except Exception as e:
            self.logger.warning(f"PyMuPDF metadata extraction error: {str(e)}")
        
        return pdf_metadata
    
    def _extract_docx_metadata(self, file_path: str) -> Dict[str, Any]:
        """
        Extract DOCX-specific metadata.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with DOCX metadata
        """
        docx_metadata = {}
        
        try:
            doc = docx.Document(file_path)
            
            # Get core properties
            core_props = doc.core_properties
            if core_props:
                if core_props.author:
                    docx_metadata['author'] = core_props.author
                if core_props.title:
                    docx_metadata['title'] = core_props.title
                if core_props.subject:
                    docx_metadata['subject'] = core_props.subject
            
            # Count document elements
            docx_metadata['paragraph_count'] = len(doc.paragraphs)
            docx_metadata['table_count'] = len(doc.tables)
            docx_metadata['page_count'] = self._estimate_page_count(doc)
            
            # Find styles used
            styles_used = set()
            for paragraph in doc.paragraphs:
                if paragraph.style:
                    styles_used.add(paragraph.style.name)
            
            docx_metadata['styles_used'] = str(list(styles_used))
            
            # Count images (a bit complex in docx)
            image_count = 0
            for rel in doc.part.rels.values():
                if rel.reltype.endswith('/image'):
                    image_count += 1
            
            docx_metadata['image_count'] = image_count
            
            # Count sections
            docx_metadata['section_count'] = len(doc.sections)
            
            # Check for headers and footers
            has_headers = False
            has_footers = False
            for section in doc.sections:
                if section.header.is_linked_to_previous == False and section.header.paragraphs:
                    has_headers = True
                if section.footer.is_linked_to_previous == False and section.footer.paragraphs:
                    has_footers = True
                    
            docx_metadata['has_headers'] = 1 if has_headers else 0
            docx_metadata['has_footers'] = 1 if has_footers else 0
            
        except Exception as e:
            self.logger.error(f"Error extracting DOCX metadata: {str(e)}")
            
        return docx_metadata
    
    def _estimate_page_count(self, doc) -> int:
        """
        Estimate page count for a DOCX document (approximate).
        
        Args:
            doc: docx.Document object
            
        Returns:
            Estimated page count
        """
        # This is a rough estimation - accurate page count requires rendering
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        
        # Rough heuristic: ~3000 chars per page, ~40 paragraphs per page, ~3 tables per page
        chars_pages = total_chars / 3000
        para_pages = total_paragraphs / 40
        table_pages = total_tables / 3
        
        # Take the max of these estimates as a conservative approach
        estimated_pages = max(chars_pages, para_pages, table_pages, 1)
        return round(estimated_pages)
    
    def _format_file_size(self, size_bytes: int) -> str:
        """
        Format file size in human-readable form.
        
        Args:
            size_bytes: File size in bytes
            
        Returns:
            Formatted file size string (e.g., "2.5 MB")
        """
        for unit in ['B', 'KB', 'MB', 'GB', 'TB']:
            if size_bytes < 1024.0 or unit == 'TB':
                break
            size_bytes /= 1024.0
        return f"{size_bytes:.2f} {unit}"
    
    def _setup_logger(self, level: str) -> logging.Logger:
        """Set up a logger instance."""
        logger = logging.getLogger(f'{__name__}.DocumentMetadataExtractor')
        
        if level.upper() == 'DEBUG':
            logger.setLevel(logging.DEBUG)
        elif level.upper() == 'INFO':
            logger.setLevel(logging.INFO)
        elif level.upper() == 'WARNING':
            logger.setLevel(logging.WARNING)
        elif level.upper() == 'ERROR':
            logger.setLevel(logging.ERROR)
        else:
            logger.setLevel(logging.INFO)
        
        # Add console handler if not already added
        if not logger.handlers:
            console_handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)
        
        return logger