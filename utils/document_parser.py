import os
import re
import logging
from typing import List, Dict, Any, Tuple, Optional, Set
import docx
import pdfplumber
from dataclasses import dataclass
import time

@dataclass
class Paragraph:
    """Class for storing paragraph information."""
    content: str
    doc_id: int
    paragraph_type: str  # 'normal', 'header', 'list', 'table', 'footer', etc.
    position: int
    header_content: Optional[str] = None

class DocumentParser:
    """
    Handles parsing of different document types (PDF, DOCX) and coordinates 
    extraction of paragraphs using appropriate methods.
    """
    def __init__(self, logging_level: str = 'INFO'):
        """Initialize the document parser."""
        self.logger = self._setup_logger(logging_level)
        self.paragraph_extractor = ParagraphExtractor(logging_level)
    
    def parse_document(self, file_path: str, doc_id: int) -> List[Paragraph]:
        """
        Parse a document and extract paragraphs.
        
        Args:
            file_path: Path to the document file
            doc_id: ID of the document in the database
            
        Returns:
            List of extracted paragraphs
        """
        self.logger.info(f"Starting to parse document: {file_path}")
        file_ext = os.path.splitext(file_path)[1].lower()
        
        start_time = time.time()
        try:
            if file_ext == '.pdf':
                paragraphs = self._parse_pdf(file_path, doc_id)
            elif file_ext in ['.docx', '.doc']:
                paragraphs = self._parse_docx(file_path, doc_id)
            else:
                self.logger.error(f"Unsupported file type: {file_ext}")
                return []
            
            # Log processing time and paragraph count
            elapsed_time = time.time() - start_time
            self.logger.info(f"Parsed {len(paragraphs)} paragraphs from {os.path.basename(file_path)} in {elapsed_time:.2f} seconds")
            
            return paragraphs
        except Exception as e:
            self.logger.error(f"Error parsing document {file_path}: {str(e)}", exc_info=True)
            return []
    
    def _parse_pdf(self, file_path: str, doc_id: int) -> List[Paragraph]:
        """
        Parse PDF document and extract paragraphs with enhanced layout analysis.
        
        Args:
            file_path: Path to the PDF file
            doc_id: ID of the document in the database
            
        Returns:
            List of extracted paragraphs
        """
        self.logger.info(f"Parsing PDF document: {file_path}")
        paragraphs = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                # Extract text from PDF with layout information
                raw_paragraphs = []
                
                for i, page in enumerate(pdf.pages):
                    self.logger.debug(f"Processing page {i+1}/{len(pdf.pages)}")
                    
                    # Extract text, tables, and layout info
                    text = page.extract_text()
                    tables = page.extract_tables()
                    
                    # Process tables
                    if tables:
                        for table in tables:
                            table_content = self._format_table(table)
                            raw_paragraphs.append({
                                'content': table_content,
                                'type': 'table',
                                'page': i+1
                            })
                    
                    # Process text (excluding tables)
                    if text:
                        # Use text layout to separate paragraphs
                        page_paragraphs = self._extract_paragraphs_from_pdf_text(text, page)
                        
                        # Log detailed information about paragraph extraction
                        self.logger.debug(f"Page {i+1}: Extracted {len(page_paragraphs)} paragraphs " 
                                         f"from {len(text)} characters of text")
                        
                        for para in page_paragraphs:
                            raw_paragraphs.append({
                                'content': para,
                                'type': 'unknown',  # Will classify later
                                'page': i+1
                            })
                
                # Process extracted content with paragraph extractor
                paragraphs = self.paragraph_extractor.process_raw_paragraphs(raw_paragraphs, doc_id)
                
                # Quality check - if we got very few paragraphs for a large document, try harder
                if len(paragraphs) <= 3 and len(raw_paragraphs) > 0:
                    total_text = sum(len(p['content']) for p in raw_paragraphs)
                    if total_text > 2000:  # Long document but few paragraphs detected
                        self.logger.warning(f"Few paragraphs detected ({len(paragraphs)}) for large document "
                                           f"({total_text} chars). Applying fallback extraction.")
                        
                        # Apply fallback extraction directly on raw content
                        enhanced_paragraphs = []
                        for raw_para in raw_paragraphs:
                            if len(raw_para['content']) > 1000:  # Only process long paragraphs
                                additional_paras = self._split_long_text(raw_para['content'])
                                for p in additional_paras:
                                    enhanced_paragraphs.append({
                                        'content': p,
                                        'type': 'unknown',
                                        'page': raw_para['page']
                                    })
                            else:
                                enhanced_paragraphs.append(raw_para)
                        
                        # Process enhanced paragraphs
                        if len(enhanced_paragraphs) > len(raw_paragraphs):
                            self.logger.info(f"Enhanced extraction: {len(raw_paragraphs)} → {len(enhanced_paragraphs)} paragraphs")
                            paragraphs = self.paragraph_extractor.process_raw_paragraphs(enhanced_paragraphs, doc_id)
                
                self.logger.info(f"Extracted {len(paragraphs)} paragraphs from PDF document")
                
            # Double-check final result
            if len(paragraphs) <= 1 and os.path.getsize(file_path) > 50000:  # Reasonable size PDF
                self.logger.warning(f"Parser only extracted {len(paragraphs)} paragraphs from a {os.path.getsize(file_path)/1024:.1f}KB PDF. Possible parsing issue.")
                    
        except Exception as e:
            self.logger.error(f"Error parsing PDF {file_path}: {str(e)}", exc_info=True)
        
        return paragraphs
    
    def _parse_docx(self, file_path: str, doc_id: int) -> List[Paragraph]:
        """
        Parse DOCX document and extract paragraphs.
        
        Args:
            file_path: Path to the DOCX file
            doc_id: ID of the document in the database
            
        Returns:
            List of extracted paragraphs
        """
        self.logger.info(f"Parsing DOCX document: {file_path}")
        paragraphs = []
        
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs, tables and lists
            raw_paragraphs = []
            position = 0
            
            # Process document elements
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    # Convert element to paragraph object
                    p = docx.text.paragraph.Paragraph(element, doc)
                    text = p.text.strip()
                    
                    if text:
                        raw_paragraphs.append({
                            'content': text,
                            'type': 'unknown',  # Will classify later
                            'style': p.style.name if hasattr(p, 'style') and p.style else 'Normal',
                            'position': position
                        })
                        position += 1
                        
                elif element.tag.endswith('tbl'):  # Table
                    table = docx.table.Table(element, doc)
                    table_content = self._extract_table_from_docx(table)
                    
                    if table_content:
                        raw_paragraphs.append({
                            'content': table_content,
                            'type': 'table',
                            'position': position
                        })
                        position += 1
            
            # Process extracted content with paragraph extractor
            paragraphs = self.paragraph_extractor.process_raw_paragraphs(raw_paragraphs, doc_id)
            
            # Check if we got reasonable content
            if len(paragraphs) == 0 and len(doc.paragraphs) > 0:
                self.logger.warning(f"No paragraphs extracted from {file_path} despite having content. Trying fallback method.")
                
                # Fallback to simpler extraction
                simple_paragraphs = []
                for i, para in enumerate(doc.paragraphs):
                    text = para.text.strip()
                    if text:
                        simple_paragraphs.append({
                            'content': text,
                            'type': 'unknown',
                            'style': para.style.name if para.style else 'Normal',
                            'position': i
                        })
                
                if simple_paragraphs:
                    paragraphs = self.paragraph_extractor.process_raw_paragraphs(simple_paragraphs, doc_id)
            
        except Exception as e:
            self.logger.error(f"Error parsing DOCX {file_path}: {str(e)}", exc_info=True)
        
        return paragraphs
    
    def _extract_paragraphs_from_pdf_text(self, text: str, page_obj=None) -> List[str]:
        """
        Extract paragraphs from PDF text using enhanced heuristics.
        
        Args:
            text: The extracted text from a PDF page
            page_obj: The pdfplumber page object (if available) for layout analysis
            
        Returns:
            List of paragraph texts
        """
        # Method 1: Split by double newlines (existing approach but improved)
        paragraphs = []
        
        # First try splitting by standard paragraph markers (double newlines)
        initial_parts = re.split(r'\n\s*\n', text)
        
        # If we only got one part, try more aggressive splitting
        if len(initial_parts) <= 1 and len(text) > 500:  # Long text with no clear paragraphs
            # Method 2: Try splitting by single newlines with additional checks
            lines = text.split('\n')
            current_para = []
            
            for i, line in enumerate(lines):
                line = line.strip()
                if not line:  # Skip empty lines
                    continue
                    
                # Start a new paragraph if:
                # 1. Current line starts with a capital letter and previous line ends with period
                # 2. Current line is indented (starts with spaces) differently than previous
                # 3. Current line starts with a bullet point or number
                # 4. Current line has different capitalization pattern (ALL CAPS vs normal)
                
                start_new_para = False
                
                # Check for sentence boundary
                if current_para and line and line[0].isupper():
                    prev_line = current_para[-1]
                    if prev_line and prev_line[-1] in '.!?':
                        # Check if this isn't a continuation (like "Dr. Smith")
                        if not (len(prev_line) >= 2 and prev_line[-2].islower() and 
                               prev_line[-1] == '.' and len(line.split()) > 3):
                            start_new_para = True
                
                # Check for bullet points or numbered lists
                if re.match(r'^\s*[\•\-\*\+○◦➢➣➤►▶→➥➔❖]\s+', line) or \
                   re.match(r'^\s*(\d+[\.\)]\s+|\([a-z\d]\)\s+|[a-z\d][\.\)]\s+)', line):
                    start_new_para = True
                
                # Check for header patterns (ALL CAPS or Title Case with no punctuation)
                if line.isupper() or (line.istitle() and not any(c in line for c in '.,:;!?')):
                    start_new_para = True
                    
                # If we detect a new paragraph and have content, save the current one
                if start_new_para and current_para:
                    paragraphs.append(' '.join(current_para))
                    current_para = []
                
                current_para.append(line)
            
            # Add the last paragraph if there's content
            if current_para:
                paragraphs.append(' '.join(current_para))
        else:
            # Process the paragraphs from the initial split
            for part in initial_parts:
                clean_part = part.strip()
                if clean_part:
                    # Further split very long paragraphs that might be incorrectly joined
                    if len(clean_part) > 1000:  # Very long paragraph
                        sub_parts = self._split_long_text(clean_part)
                        paragraphs.extend(sub_parts)
                    else:
                        paragraphs.append(clean_part)
        
        # Method 3: Use layout analysis if available and we still don't have good paragraphs
        if page_obj and (len(paragraphs) <= 1) and len(text) > 500:
            try:
                # Extract character data with positions
                chars = page_obj.chars
                if chars:
                    layout_paragraphs = self._extract_paragraphs_from_layout(chars)
                    if len(layout_paragraphs) > len(paragraphs):
                        return layout_paragraphs
            except Exception as e:
                self.logger.warning(f"Layout analysis failed: {str(e)}")
        
        # If we still couldn't parse paragraphs, use a fallback method
        if len(paragraphs) <= 1 and len(text) > 500:
            # Fallback: Split by sentences as a last resort
            sentences = re.split(r'(?<=[.!?])\s+(?=[A-Z])', text)
            # Group sentences into paragraphs (roughly 3-5 sentences per paragraph)
            if len(sentences) > 3:
                grouped_sentences = []
                for i in range(0, len(sentences), 4):  # Group by ~4 sentences
                    group = ' '.join(sentences[i:i+4])
                    if group.strip():
                        grouped_sentences.append(group)
                if grouped_sentences:
                    return grouped_sentences
        
        return paragraphs

    def _split_long_text(self, text: str) -> List[str]:
        """
        Split very long text into paragraph-like chunks based on sentence boundaries.
        
        Args:
            text: Long text to split
            
        Returns:
            List of text chunks
        """
        # Find sentence boundaries
        sentence_boundaries = [m.end() for m in re.finditer(r'[.!?]\s+', text)]
        
        if not sentence_boundaries:
            return [text]  # No clear sentences found
        
        # If average sentence length is very long, this might not be proper paragraphs
        avg_sentence_length = sentence_boundaries[0] if len(sentence_boundaries) == 1 else \
                            sentence_boundaries[-1] / len(sentence_boundaries)
        
        # Group sentences into paragraph-like chunks
        chunks = []
        start = 0
        
        # Aim for paragraphs of 3-5 sentences or ~500 characters
        target_chunk_size = min(3 if avg_sentence_length > 100 else 5, 
                            max(1, int(500 / avg_sentence_length)))
        
        for i in range(target_chunk_size, len(sentence_boundaries), target_chunk_size):
            end = sentence_boundaries[min(i, len(sentence_boundaries) - 1)]
            chunks.append(text[start:end].strip())
            start = end
        
        # Add the last chunk
        if start < len(text):
            chunks.append(text[start:].strip())
        
        return chunks

    def _extract_paragraphs_from_layout(self, chars: List[Dict]) -> List[str]:
        """
        Extract paragraphs based on character layout analysis.
        
        Args:
            chars: List of character data from pdfplumber
            
        Returns:
            List of extracted paragraphs
        """
        if not chars:
            return []
        
        # Sort characters by y0 (vertical position) then x0 (horizontal position)
        sorted_chars = sorted(chars, key=lambda c: (c['top'], c['x0']))
        
        # Group characters into lines based on y-position
        lines = []
        current_line = []
        current_y = None
        
        for char in sorted_chars:
            if current_y is None or abs(char['top'] - current_y) < 2:  # Same line (with small tolerance)
                current_line.append(char)
                current_y = char['top']
            else:  # New line
                if current_line:
                    lines.append(current_line)
                current_line = [char]
                current_y = char['top']
        
        # Add the last line
        if current_line:
            lines.append(current_line)
        
        # Convert lines to text
        text_lines = []
        for line in lines:
            # Sort by x position within each line
            sorted_line = sorted(line, key=lambda c: c['x0'])
            line_text = ''.join(c['text'] for c in sorted_line if 'text' in c)
            if line_text.strip():
                text_lines.append(line_text.strip())
        
        # Group lines into paragraphs based on vertical spacing
        paragraphs = []
        current_para = []
        prev_bottom = None
        
        for i, line in enumerate(text_lines):
            if not line.strip():
                continue
                
            # If we have line position data
            if i < len(lines) and lines[i]:
                current_top = min(c['top'] for c in lines[i])
                
                # Check if this is a new paragraph based on spacing
                if prev_bottom is not None:
                    # Estimate line height
                    line_height = max(c['height'] for c in lines[i] if 'height' in c) \
                                if any('height' in c for c in lines[i]) else 10
                    
                    # If the gap is significantly larger than a typical line height, start new paragraph
                    if (current_top - prev_bottom) > (line_height * 1.5):
                        if current_para:
                            paragraphs.append(' '.join(current_para))
                            current_para = []
                
                # Update for next iteration
                prev_bottom = max(c['bottom'] for c in lines[i] if 'bottom' in c) \
                            if any('bottom' in c for c in lines[i]) else current_top + 10
            
            # Apply text-based heuristics as well
            if current_para and (line.isupper() or re.match(r'^\d+[\.\)]\s+', line)):
                # Headers or numbered items should start new paragraphs
                if current_para:
                    paragraphs.append(' '.join(current_para))
                    current_para = []
            
            current_para.append(line)
        
        # Add the last paragraph
        if current_para:
            paragraphs.append(' '.join(current_para))
        
        return paragraphs
    
    def _format_table(self, table: List[List[str]]) -> str:
        """
        Format a pdfplumber table into a string representation.
        
        Args:
            table: Table data from pdfplumber
            
        Returns:
            Formatted table string
        """
        # Filter out empty/None cells and replace with empty string
        formatted_table = [[cell if cell else '' for cell in row] for row in table]
        
        # Convert to string representation
        result = []
        for row in formatted_table:
            result.append(" | ".join(str(cell).strip() for cell in row))
        
        return "\n".join(result)
    
    def _extract_table_from_docx(self, table) -> str:
        """
        Extract content from a docx table.
        
        Args:
            table: DOCX table object
            
        Returns:
            Formatted table string
        """
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cells.append(cell.text.strip())
            rows.append(" | ".join(cells))
        
        return "\n".join(rows)
    
    def _setup_logger(self, level: str) -> logging.Logger:
        """
        Set up a logger instance.
        
        Args:
            level: Logging level
            
        Returns:
            Configured logger
        """
        logger = logging.getLogger(f'{__name__}.DocumentParser')
        
        if not logger.handlers:  # Only add handlers if they don't exist
            # Set level
            if level.upper() == 'DEBUG':
                logger.setLevel(logging.DEBUG)
            elif level.upper() == 'INFO':
                logger.setLevel(logging.INFO)
            elif level.upper() == 'WARNING':
                logger.setLevel(logging.WARNING)
            elif level.upper() == 'ERROR':
                logger.setLevel(logging.ERROR)
            else:
                logger.setLevel(logging.INFO)
            
            # Create console handler
            console_handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)
        
        return logger


class ParagraphExtractor:
    """
    Handles the extraction and classification of paragraphs from document content.
    Implements multiple strategies for paragraph detection.
    """
    def __init__(self, logging_level: str = 'INFO'):
        """Initialize the paragraph extractor."""
        self.logger = self._setup_logger(logging_level)
        self.stopwords = self._get_stopwords()
    
    def _get_stopwords(self) -> Set[str]:
        """Get a set of common stopwords for text processing."""
        stopwords = {
            'a', 'an', 'the', 'and', 'or', 'but', 'if', 'then', 'else', 'when',
            'at', 'by', 'for', 'with', 'about', 'against', 'between', 'into',
            'to', 'from', 'in', 'out', 'on', 'off', 'over', 'under', 'again',
            'further', 'then', 'once', 'here', 'there', 'all', 'any', 'both',
            'each', 'few', 'more', 'most', 'other', 'some', 'such', 'no', 'nor',
            'not', 'only', 'own', 'same', 'so', 'than', 'too', 'very', 'can',
            'will', 'just', 'should', 'now', 'is', 'are', 'was', 'were', 'be',
            'been', 'being', 'have', 'has', 'had', 'having', 'do', 'does', 'did',
            'doing', 'this', 'that', 'these', 'those', 'of', 'up', 'down'
        }
        return stopwords
    
    def process_raw_paragraphs(self, raw_paragraphs: List[Dict], doc_id: int) -> List[Paragraph]:
        """
        Process raw paragraphs and return structured Paragraph objects.
        
        Args:
            raw_paragraphs: List of raw paragraph data
            doc_id: Document ID
            
        Returns:
            List of structured Paragraph objects
        """
        self.logger.info(f"Processing {len(raw_paragraphs)} raw paragraphs for document {doc_id}")
        
        # 1. Classify paragraph types
        classified_paragraphs = self._classify_paragraphs(raw_paragraphs)
        
        # 2. Combine lists
        combined_lists = self._combine_lists(classified_paragraphs)
        
        # 3. Associate headers with paragraphs
        with_headers = self._associate_headers(combined_lists)
        
        # 4. Create final paragraph objects
        paragraphs = []
        for i, para in enumerate(with_headers):
            header_content = para.get('header_content')
            
            # Skip if this is an address or empty paragraph
            if para['type'] == 'address' or not para['content'].strip():
                self.logger.debug(f"Skipping {para['type']} paragraph: {para['content'][:30]}...")
                continue
                
            paragraph = Paragraph(
                content=para['content'],
                doc_id=doc_id,
                paragraph_type=para['type'],
                position=i,
                header_content=header_content
            )
            paragraphs.append(paragraph)
        
        self.logger.info(f"Created {len(paragraphs)} structured paragraphs")
        return paragraphs
    
    def _classify_paragraphs(self, raw_paragraphs: List[Dict]) -> List[Dict]:
        """
        Classify paragraphs into different types.
        
        Args:
            raw_paragraphs: List of raw paragraph data
            
        Returns:
            List of classified paragraphs
        """
        classified = []
        
        for para in raw_paragraphs:
            para_type = para.get('type', 'unknown')
            content = para.get('content', '').strip()
            
            # Skip empty paragraphs
            if not content:
                continue
            
            # If already classified as table, keep it
            if para_type == 'table':
                classified.append(para)
                continue
            
            # Check if paragraph is a header
            if self._is_header(content, para):
                para['type'] = 'header'
                classified.append(para)
                continue
            
            # Check if paragraph is a list
            if self._is_list(content):
                para['type'] = 'list'
                classified.append(para)
                continue
            
            # Check if paragraph is an address
            if self._is_address(content):
                para['type'] = 'address'
                classified.append(para)
                continue
            
            # Check if paragraph is a footer/disclaimer
            if self._is_boilerplate(content):
                para['type'] = 'boilerplate'
                classified.append(para)
                continue
            
            # Default to normal paragraph
            para['type'] = 'normal'
            classified.append(para)
        
        return classified
    
    def _combine_lists(self, paragraphs: List[Dict]) -> List[Dict]:
        """
        Combine consecutive list items into a single paragraph.
        
        Args:
            paragraphs: List of paragraphs
            
        Returns:
            List of paragraphs with combined lists
        """
        if not paragraphs:
            return []
            
        combined = []
        current_list = None
        
        for para in paragraphs:
            if para['type'] == 'list':
                if current_list is None:
                    current_list = {
                        'content': para['content'],
                        'type': 'list',
                        'position': para.get('position', 0)
                    }
                else:
                    current_list['content'] += '\n' + para['content']
            else:
                if current_list is not None:
                    combined.append(current_list)
                    current_list = None
                combined.append(para)
        
        # Add the last list if there is one
        if current_list is not None:
            combined.append(current_list)
        
        return combined
    
    def _associate_headers(self, paragraphs: List[Dict]) -> List[Dict]:
        """
        Associate headers with their following paragraphs.
        
        Args:
            paragraphs: List of paragraphs
            
        Returns:
            List of paragraphs with header associations
        """
        if not paragraphs:
            return []
            
        result = []
        last_header = None
        
        for para in paragraphs:
            if para['type'] == 'header':
                last_header = para['content']
                # Add the header as a standalone paragraph too
                result.append(para)
            else:
                if last_header is not None:
                    para['header_content'] = last_header
                    last_header = None  # Only associate with the immediately following paragraph
                result.append(para)
        
        return result
    
    def _is_header(self, content: str, para_info: Dict) -> bool:
        """
        Determine if a paragraph is a header.
        
        Args:
            content: Paragraph content
            para_info: Additional paragraph information
            
        Returns:
            True if paragraph is a header, False otherwise
        """
        # Check if style indicates a header (for DOCX)
        if 'style' in para_info and any(term in para_info['style'].lower() for term in ['head', 'title', 'subtitle']):
            return True
            
        # For other cases, use heuristics
        # Headers are typically short
        if len(content) > 100:
            return False
            
        # Headers often end without punctuation
        if not content[-1] in '.?!:;,':
            # Check for common header patterns
            if re.match(r'^[\d\.]+\s+\w+', content):  # Numbered header (e.g., "1.2 Introduction")
                return True
            if content.isupper():  # ALL CAPS headers
                return True
            if content.istitle() and len(content.split()) <= 10:  # Title Case headers
                return True
            if len(content.split()) <= 6:  # Short text without punctuation is likely a header
                return True
        
        return False
    
    def _is_list(self, content: str) -> bool:
        """
        Determine if a paragraph is a list item.
        
        Args:
            content: Paragraph content
            
        Returns:
            True if paragraph is a list item, False otherwise
        """
        # Check for bullet points
        if re.match(r'^\s*[•\-\*\+○◦➢➣➤►▶→➥➔❖]\s+', content):
            return True
            
        # Check for numbered lists
        if re.match(r'^\s*(\d+[\.\)]\s+|\([a-z\d]\)\s+|[a-z\d][\.\)]\s+)', content):
            return True
            
        return False
    
    def _is_address(self, content: str) -> bool:
        """
        Determine if a paragraph is an address.
        
        Args:
            content: Paragraph content
            
        Returns:
            True if paragraph is an address, False otherwise
        """
        # Check for postal code patterns
        postal_pattern = re.search(r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2}\b', content)  # UK
        zip_pattern = re.search(r'\b\d{5}(-\d{4})?\b', content)  # US
        
        # Check for typical address components
        address_indicators = [
            'street', 'avenue', 'road', 'lane', 'drive', 'boulevard', 
            'st.', 'ave.', 'rd.', 'ln.', 'dr.', 'blvd.', 'suite', 'apt', 
            'apartment', 'floor', 'unit'
        ]
        has_indicator = any(indicator in content.lower() for indicator in address_indicators)
        
        # If we find postal codes or address indicators, it's likely an address
        if (postal_pattern or zip_pattern) and has_indicator:
            return True
            
        return False
    
    def _is_boilerplate(self, content: str) -> bool:
        """
        Determine if a paragraph is boilerplate text (footer, disclaimer, etc.).
        
        Args:
            content: Paragraph content
            
        Returns:
            True if paragraph is boilerplate text, False otherwise
        """
        # Check for common boilerplate indicators
        boilerplate_indicators = [
            'all rights reserved', 'copyright', '©', 'confidential', 
            'disclaimer', 'terms and conditions', 'privacy policy',
            'legal notice', 'proprietary', 'confidentiality notice',
            'do not copy', 'not for distribution', 'all rights reserved'
        ]
        
        if any(indicator in content.lower() for indicator in boilerplate_indicators):
            return True
            
        # Check for typical footer patterns (page numbers, etc.)
        if re.match(r'^\s*Page \d+ of \d+\s*$', content):
            return True
            
        return False
    
    def _setup_logger(self, level: str) -> logging.Logger:
        """
        Set up a logger instance.
        
        Args:
            level: Logging level
            
        Returns:
            Configured logger
        """
        logger = logging.getLogger(f'{__name__}.ParagraphExtractor')
        
        if not logger.handlers:  # Only add handlers if they don't exist
            # Set level
            if level.upper() == 'DEBUG':
                logger.setLevel(logging.DEBUG)
            elif level.upper() == 'INFO':
                logger.setLevel(logging.INFO)
            elif level.upper() == 'WARNING':
                logger.setLevel(logging.WARNING)
            elif level.upper() == 'ERROR':
                logger.setLevel(logging.ERROR)
            else:
                logger.setLevel(logging.INFO)
            
            # Create console handler
            console_handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)
        
        return logger