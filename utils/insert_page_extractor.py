"""
Insert Page Extractor module for extracting full page content from inserts.
"""

import logging
import os
from typing import List, Dict, Any, Optional

import fitz  # PyMuPDF
import docx

class InsertPageExtractor:
    """
    Extracts full page content from insert documents for matching.
    Keeps pages as whole units rather than breaking into paragraphs.
    """
    
    def __init__(self, logging_level: str = 'INFO'):
        """Initialize the page extractor."""
        self.logger = self._setup_logger(logging_level)
    
    def extract_pages(self, file_path: str, insert_id: int) -> List[Dict[str, Any]]:
        """
        Extract pages from a document file.
        
        Args:
            file_path: Path to the document file
            insert_id: ID of the insert in the database
            
        Returns:
            List of dictionaries with page content and metadata
        """
        self.logger.info(f"Extracting pages from: {file_path}")
        
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_pages_from_pdf(file_path, insert_id)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_pages_from_docx(file_path, insert_id)
        else:
            self.logger.error(f"Unsupported file type: {file_ext}")
            return []
    
    def _extract_pages_from_pdf(self, file_path: str, insert_id: int) -> List[Dict[str, Any]]:
        """
        Extract pages from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            insert_id: ID of the insert in the database
            
        Returns:
            List of dictionaries with page content and metadata
        """
        pages = []
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                # Extract all text from the page
                text = page.get_text()
                
                if text.strip():  # Only add non-empty pages
                    pages.append({
                        'content': text,
                        'page_number': page_num,
                        'insert_id': insert_id
                    })
            
            self.logger.info(f"Extracted {len(pages)} pages from PDF")
            
        except Exception as e:
            self.logger.error(f"Error extracting pages from PDF: {str(e)}", exc_info=True)
        
        return pages
    
    def _extract_pages_from_docx(self, file_path: str, insert_id: int) -> List[Dict[str, Any]]:
        """
        Extract pages from a DOCX file.
        Note: Page boundaries in DOCX are estimated since the format doesn't explicitly define pages.
        
        Args:
            file_path: Path to the DOCX file
            insert_id: ID of the insert in the database
            
        Returns:
            List of dictionaries with page content and metadata
        """
        pages = []
        
        try:
            doc = docx.Document(file_path)
            
            # Collect all paragraphs
            all_paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    all_paragraphs.append(para.text)
            
            # Collect all tables as text
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_text.append(' | '.join(cells))
                if table_text:
                    all_paragraphs.append('\n'.join(table_text))
            
            # Estimate page boundaries (rough estimate - about 3000 characters per page)
            CHARS_PER_PAGE = 3000
            current_page = []
            current_chars = 0
            page_num = 0
            
            for para in all_paragraphs:
                if current_chars + len(para) > CHARS_PER_PAGE and current_page:
                    # Create a new page
                    pages.append({
                        'content': '\n'.join(current_page),
                        'page_number': page_num,
                        'insert_id': insert_id
                    })
                    current_page = [para]
                    current_chars = len(para)
                    page_num += 1
                else:
                    current_page.append(para)
                    current_chars += len(para)
            
            # Add the last page if there's content
            if current_page:
                pages.append({
                    'content': '\n'.join(current_page),
                    'page_number': page_num,
                    'insert_id': insert_id
                })
            
            self.logger.info(f"Extracted {len(pages)} estimated pages from DOCX")
            
        except Exception as e:
            self.logger.error(f"Error extracting pages from DOCX: {str(e)}", exc_info=True)
        
        return pages
    
    def _setup_logger(self, level: str) -> logging.Logger:
        """Set up a logger instance."""
        logger = logging.getLogger(f'{__name__}.InsertPageExtractor')
        
        if level.upper() == 'DEBUG':
            logger.setLevel(logging.DEBUG)
        elif level.upper() == 'INFO':
            logger.setLevel(logging.INFO)
        elif level.upper() == 'WARNING':
            logger.setLevel(logging.WARNING)
        elif level.upper() == 'ERROR':
            logger.setLevel(logging.ERROR)
        else:
            logger.setLevel(logging.INFO)
        
        # Add console handler if not already added
        if not logger.handlers:
            console_handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)
        
        return logger