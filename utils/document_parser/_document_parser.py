import os
import re
import logging
import json
from typing import List, Dict, Any, Tuple, Optional, Set, Union
from dataclasses import dataclass, field, asdict
import time
from collections import Counter
import numpy as np
from itertools import groupby
from operator import itemgetter

# PDF libraries
try:
    import fitz  # PyMuPDF
except ImportError:
    print("Warning: PyMuPDF not installed. PDF parsing may be limited.")
    fitz = None

try:
    import pdfplumber
except ImportError:
    print("Warning: PDFPlumber not installed. Table extraction from PDFs may be limited.")
    pdfplumber = None

# DOCX library
try:
    import docx
except ImportError:
    print("Warning: python-docx not installed. DOCX parsing will not be available.")
    docx = None

# Optional OCR library
try:
    import pytesseract
    from PIL import Image
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


# Original Paragraph class to maintain compatibility with database_manager.py
@dataclass
class Paragraph:
    """Class for storing paragraph information."""
    content: str
    doc_id: int
    paragraph_type: str  # 'normal', 'header', 'list', 'table', 'footer', etc.
    position: int
    header_content: Optional[str] = None


@dataclass
class DocumentElement:
    """Base class for document elements."""
    content: str
    element_type: str  # 'paragraph', 'heading', 'list', 'list_item', 'table', 'toc', etc.
    page_num: int = 0
    position: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert element to dictionary."""
        result = asdict(self)
        return result
    
    def to_paragraph(self, doc_id: int) -> Paragraph:
        """Convert DocumentElement to Paragraph for compatibility."""
        # Map element_type to paragraph_type
        type_mapping = {
            'paragraph': 'normal',
            'heading': 'header',
            'list': 'list',
            'table': 'table',
            'toc': 'normal',
            'toc_entry': 'normal',
            'boilerplate': 'boilerplate',
            'address': 'address',
            'footer': 'footer',
        }
        
        # Get header content from metadata if available
        header_content = None
        if 'heading' in self.metadata:
            header_content = self.metadata['heading']
        
        return Paragraph(
            content=self.content,
            doc_id=doc_id,
            paragraph_type=type_mapping.get(self.element_type, 'normal'),
            position=self.position,
            header_content=header_content
        )


@dataclass
class Document:
    """Represents a parsed document with its elements."""
    filename: str
    elements: List[DocumentElement] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert document to dictionary."""
        return {
            'filename': self.filename,
            'metadata': self.metadata,
            'elements': [element.to_dict() for element in self.elements]
        }


class DocumentParser:
    """
    Main parser for document extraction.
    
    Handles parsing of different document types (PDF, DOCX) and coordinates 
    extraction of text, lists, tables, and other structural elements.
    """
    def __init__(self, ocr_enabled: bool = False, logging_level: str = 'INFO'):
        """Initialize document parser."""
        self.ocr_enabled = ocr_enabled and HAS_OCR
        self.logger = self._setup_logger(logging_level)
    
    def parse_document(self, file_path: str, doc_id: int) -> List[Paragraph]:
        """
        Parse a document and extract paragraphs.
        
        Args:
            file_path: Path to the document file
            doc_id: ID of the document in the database
            
        Returns:
            List of extracted paragraphs
        """
        self.logger.info(f"Starting to parse document: {file_path}")
        file_ext = os.path.splitext(file_path)[1].lower()
        
        start_time = time.time()
        try:
            # Use the new parsing logic but return Paragraph objects
            document = self._parse(file_path)
            
            # Convert DocumentElement objects to Paragraph objects
            paragraphs = [element.to_paragraph(doc_id) for element in document.elements]
            
            # Post-process paragraphs
            paragraphs = self._post_process_paragraphs(paragraphs)
            
            # Log processing time
            elapsed_time = time.time() - start_time
            self.logger.info(f"Parsed {len(paragraphs)} paragraphs from {os.path.basename(file_path)} in {elapsed_time:.2f} seconds")
            
            return paragraphs
            
        except Exception as e:
            self.logger.error(f"Error parsing document {file_path}: {str(e)}", exc_info=True)
            return []
    
    def _parse(self, file_path: str) -> Document:
        """
        Parse a document and extract its content.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object with extracted elements
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        elements = []
        
        if file_ext == '.pdf':
            if fitz is None or pdfplumber is None:
                self.logger.error("PDF parsing requires PyMuPDF and PDFPlumber libraries")
                return Document(filename=file_path)
            elements = self._parse_pdf_elements(file_path)
        elif file_ext in ['.docx', '.doc']:
            if docx is None:
                self.logger.error("DOCX parsing requires python-docx library")
                return Document(filename=file_path)
            elements = self._parse_docx_elements(file_path)
        else:
            self.logger.error(f"Unsupported file type: {file_ext}")
            return Document(filename=file_path)
        
        # Create document with elements
        document = Document(
            filename=file_path,
            elements=elements,
            metadata={
                'file_type': file_ext,
                'parser_version': '1.0.0',
                'parsed_at': time.strftime('%Y-%m-%d %H:%M:%S'),
                'element_count': len(elements)
            }
        )
        
        return document
    
    def _parse_pdf_elements(self, file_path: str) -> List[DocumentElement]:
        """
        Parse PDF document using a combination of PyMuPDF and PDFPlumber.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            List of extracted document elements
        """
        self.logger.info(f"Parsing PDF document: {file_path}")
        elements = []
        position = 0
        
        try:
            # First use PDFPlumber for tables to ensure proper table detection
            tables_by_page = {}
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        tables_by_page[page_num] = []
                        for table in tables:
                            if table and any(any(cell for cell in row) for row in table):
                                # Store table data and its bounding box for later use
                                table_areas = page.find_tables(table_settings={"vertical_strategy": "text", 
                                                                               "horizontal_strategy": "text"})
                                table_bbox = None
                                if table_areas and len(table_areas) > len(tables_by_page[page_num]):
                                    table_bbox = (
                                        table_areas[len(tables_by_page[page_num])].bbox[0],
                                        table_areas[len(tables_by_page[page_num])].bbox[1],
                                        table_areas[len(tables_by_page[page_num])].bbox[2],
                                        table_areas[len(tables_by_page[page_num])].bbox[3],
                                    )
                                
                                tables_by_page[page_num].append({
                                    'data': table,
                                    'bbox': table_bbox
                                })
            
            # Then use PyMuPDF for document structure and text
            doc = fitz.open(file_path)
            
            # Check for potential scan-only PDF
            is_scanned = self._is_scanned_pdf(doc)
            
            if is_scanned and self.ocr_enabled:
                self.logger.info("Detected scanned PDF, using OCR")
                elements = self._process_with_ocr(file_path)
            else:
                # Extract document structure
                toc = doc.get_toc()
                if toc:
                    toc_element = self._process_toc(toc)
                    elements.append(toc_element)
                    position += 1
                
                # Process each page
                for page_num, page in enumerate(doc):
                    # Extract and add tables first (if any) to prevent text extraction from cutting them
                    if page_num in tables_by_page:
                        for table_info in tables_by_page[page_num]:
                            table_element = DocumentElement(
                                content=self._format_table_content(self._clean_table(table_info['data'])),
                                element_type="table",
                                page_num=page_num + 1,
                                position=position,
                                metadata={
                                    'rows': self._clean_table(table_info['data']),
                                    'bbox': table_info['bbox']
                                }
                            )
                            elements.append(table_element)
                            position += 1
                    
                    # Process text content, exclude table areas
                    table_rects = []
                    if page_num in tables_by_page:
                        for table_info in tables_by_page[page_num]:
                            if table_info['bbox']:
                                # Create rectangle for table area to exclude
                                x0, y0, x1, y1 = table_info['bbox']
                                table_rects.append((x0, y0, x1, y1))
                    
                    # IMPORTANT: We'll use a much more conservative approach for column detection
                    # to avoid cutting paragraphs in the middle of sentences
                    try:
                        # First attempt to process as single column (safer default)
                        blocks = page.get_text("dict")["blocks"]
                        
                        # Filter blocks that overlap with tables
                        filtered_blocks = self._filter_blocks_by_tables(blocks, table_rects)
                        
                        # Process blocks
                        page_elements = self._process_blocks(filtered_blocks, page_num + 1)
                        
                        # Filter out headers and footers
                        page_elements = self._filter_headers_footers(page_elements, page.rect.height)
                        
                        # Only try column detection if we have clear signs of a multi-column layout
                        # Check for column layout (CONSERVATIVE approach - requires clear evidence)
                        is_multi_column = self._is_clearly_multi_column(filtered_blocks, page.rect.width)
                        
                        if is_multi_column:
                            # Use multi-column processing when we're confident
                            self.logger.info(f"Page {page_num+1} appears to have multiple columns, processing accordingly")
                            col_page_elements = self._process_multi_column_page(page, table_rects, page_num)
                            
                            # Only use multi-column results if they seem valid
                            if self._is_valid_column_segmentation(col_page_elements):
                                page_elements = col_page_elements
                    
                    except Exception as e:
                        self.logger.warning(f"Error in column analysis for page {page_num+1}, using fallback: {str(e)}")
                        # Fallback to simple text extraction
                        text = page.get_text()
                        if text.strip():
                            paragraphs = re.split(r'\n\s*\n', text)
                            
                            page_elements = []
                            for para in paragraphs:
                                para = para.strip()
                                if not para:
                                    continue
                                
                                element_type = 'paragraph'
                                if self._is_list_item(para):
                                    element_type = 'list'
                                elif len(para) < 100 and (para.isupper() or para[-1] not in '.?!:;,'):
                                    element_type = 'heading'
                                
                                page_elements.append(DocumentElement(
                                    content=para,
                                    element_type=element_type,
                                    page_num=page_num + 1,
                                    position=0,  # Will be updated later
                                    metadata={}
                                ))
                    
                    # Update positions for all elements on this page
                    for element in page_elements:
                        element.position = position
                        position += 1
                    
                    elements.extend(page_elements)
            
            doc.close()
        except Exception as e:
            self.logger.error(f"Error in PDF parsing: {str(e)}", exc_info=True)
        
        # Post-process to combine related elements
        elements = self._combine_related_elements(elements)
        
        return elements
    
    def _is_clearly_multi_column(self, blocks, page_width):
        """
        Determine if a page clearly has multiple columns.
        Uses conservative heuristics to avoid false positives.
        
        Args:
            blocks: Text blocks from PyMuPDF
            page_width: Width of the page
            
        Returns:
            Boolean indicating if page is clearly multi-column
        """
        # Need enough blocks to make a determination
        if len(blocks) < 10:
            return False
        
        # Get text blocks only
        text_blocks = [b for b in blocks if b["type"] == 0 and "lines" in b and b["lines"]]
        
        if len(text_blocks) < 5:
            return False
        
        # Get x-midpoints of all blocks
        midpoints = []
        for block in text_blocks:
            if "bbox" in block:
                x0, y0, x1, y1 = block["bbox"]
                midpoint = (x0 + x1) / 2
                midpoints.append(midpoint)
        
        if not midpoints:
            return False
        
        # Check distribution of midpoints - use histogram
        hist, bins = np.histogram(midpoints, bins=20)
        
        # Look for clear gaps in the middle area of the page
        middle_start = int(len(hist) * 0.3)  # 30% from left
        middle_end = int(len(hist) * 0.7)    # 70% from left
        middle_region = hist[middle_start:middle_end]
        
        # We need a clear gap (near-zero values) in the middle
        if len(middle_region) > 0 and np.min(middle_region) < 0.1 * np.max(hist):
            # Check for good distribution on both sides
            left_count = np.sum(hist[:middle_start])
            right_count = np.sum(hist[middle_end:])
            
            # Both left and right sides should have substantial content
            if left_count > 3 and right_count > 3:
                # Check block widths - columns typically have blocks less than 40% of page width
                narrow_blocks = 0
                for block in text_blocks:
                    if "bbox" in block:
                        x0, _, x1, _ = block["bbox"]
                        width = x1 - x0
                        if width < 0.4 * page_width:
                            narrow_blocks += 1
                
                # At least 70% of blocks should be narrow
                if narrow_blocks >= 0.7 * len(text_blocks):
                    return True
        
        return False
    
    def _process_multi_column_page(self, page, table_rects, page_num):
        """
        Process a page with multiple columns.
        
        Args:
            page: PyMuPDF page object
            table_rects: List of table rectangles to exclude
            page_num: Page number
            
        Returns:
            List of document elements
        """
        # Get page dimensions
        page_width = page.rect.width
        
        # Analyze blocks for column detection
        blocks = page.get_text("dict")["blocks"]
        filtered_blocks = self._filter_blocks_by_tables(blocks, table_rects)
        
        # Find column boundaries using more sophisticated analysis
        column_bounds = self._detect_columns(filtered_blocks, page_width)
        
        if not column_bounds or len(column_bounds) < 2:
            # Fallback to single column if no clear columns detected
            return self._process_blocks(filtered_blocks, page_num + 1)
        
        # Process each column
        elements = []
        
        for col_idx, (x0, x1) in enumerate(column_bounds):
            # Create rectangle for column
            col_rect = fitz.Rect(x0, 0, x1, page.rect.height)
            
            # Get blocks in this column
            col_blocks = page.get_text("dict", clip=col_rect)["blocks"]
            
            # Filter blocks that overlap with tables
            filtered_col_blocks = self._filter_blocks_by_tables(col_blocks, table_rects)
            
            # Process blocks in this column
            col_elements = self._process_blocks(filtered_col_blocks, page_num + 1)
            elements.extend(col_elements)
        
        return elements
    
    def _detect_columns(self, blocks, page_width):
        """
        Detect column boundaries using advanced analysis.
        
        Args:
            blocks: Text blocks from PyMuPDF
            page_width: Width of the page
            
        Returns:
            List of column bounds (x0, x1) tuples
        """
        # Get text blocks only
        text_blocks = [b for b in blocks if b["type"] == 0 and "lines" in b and b["lines"]]
        
        if len(text_blocks) < 5:
            return []
        
        # Get block boundaries
        block_bounds = []
        for block in text_blocks:
            if "bbox" in block:
                x0, y0, x1, y1 = block["bbox"]
                block_bounds.append((x0, x1))
        
        # Group blocks by horizontal position
        block_bounds.sort()  # Sort by x0
        
        # Look for gaps between blocks
        all_bounds = []
        for x0, x1 in block_bounds:
            all_bounds.append((x0, "start"))
            all_bounds.append((x1, "end"))
        
        all_bounds.sort()
        
        # Find gaps
        active_blocks = 0
        potential_gaps = []
        
        for pos, action in all_bounds:
            if action == "start":
                active_blocks += 1
            else:  # end
                active_blocks -= 1
            
            # When active_blocks is 0, we have a gap
            if active_blocks == 0:
                gap_start = pos
                # Find next start position
                next_start_idx = all_bounds.index((pos, action)) + 1
                if next_start_idx < len(all_bounds):
                    gap_end = all_bounds[next_start_idx][0]
                    gap_width = gap_end - gap_start
                    
                    # Only consider significant gaps
                    if gap_width > 0.05 * page_width:
                        potential_gaps.append((gap_start, gap_end))
        
        # If we don't have any significant gaps, this is likely not a multi-column page
        if not potential_gaps or len(potential_gaps) < 1:
            return []
        
        # Group gaps that are close together
        merged_gaps = []
        current_gap = potential_gaps[0]
        
        for i in range(1, len(potential_gaps)):
            if potential_gaps[i][0] - current_gap[1] < 0.05 * page_width:
                # Merge gaps
                current_gap = (current_gap[0], potential_gaps[i][1])
            else:
                merged_gaps.append(current_gap)
                current_gap = potential_gaps[i]
        
        merged_gaps.append(current_gap)
        
        # Create column bounds
        if not merged_gaps:
            return []
        
        column_bounds = []
        column_start = 0
        
        for gap_start, gap_end in merged_gaps:
            column_bounds.append((column_start, gap_start))
            column_start = gap_end
        
        # Add the last column
        column_bounds.append((column_start, page_width))
        
        # Verify that we have reasonable columns (not too narrow)
        valid_bounds = []
        for x0, x1 in column_bounds:
            width = x1 - x0
            if width > 0.15 * page_width:  # Column should be at least 15% of page width
                valid_bounds.append((x0, x1))
        
        return valid_bounds if len(valid_bounds) >= 2 else []
    
    def _filter_blocks_by_tables(self, blocks, table_rects):
        """
        Filter blocks that overlap with tables.
        
        Args:
            blocks: List of text blocks
            table_rects: List of table rectangles
            
        Returns:
            Filtered list of blocks
        """
        if not table_rects:
            return blocks
            
        filtered_blocks = []
        for block in blocks:
            if "bbox" in block:
                block_bbox = block["bbox"]
                # Check if block overlaps significantly with any table
                overlaps_table = False
                for table_rect in table_rects:
                    x_overlap = max(0, min(block_bbox[2], table_rect[2]) - max(block_bbox[0], table_rect[0]))
                    y_overlap = max(0, min(block_bbox[3], table_rect[3]) - max(block_bbox[1], table_rect[1]))
                    overlap_area = x_overlap * y_overlap
                    block_area = (block_bbox[2] - block_bbox[0]) * (block_bbox[3] - block_bbox[1])
                    
                    if block_area > 0 and overlap_area / block_area > 0.5:  # 50% overlap threshold
                        overlaps_table = True
                        break
                
                if not overlaps_table:
                    filtered_blocks.append(block)
            else:
                filtered_blocks.append(block)
                
        return filtered_blocks
    
    def _is_valid_column_segmentation(self, elements):
        """
        Check if column segmentation resulted in valid elements.
        
        Args:
            elements: List of document elements
            
        Returns:
            Boolean indicating if segmentation appears valid
        """
        # Too few elements may indicate bad segmentation
        if len(elements) < 3:
            return False
            
        # Check for very short paragraphs that might be fragments
        fragment_count = 0
        for element in elements:
            # Look for short text that ends without punctuation
            if (element.element_type == 'paragraph' and 
                len(element.content) < 50 and 
                element.content and
                element.content[-1] not in '.?!:;'):
                fragment_count += 1
        
        # If too many fragments, segmentation may be invalid
        return fragment_count < len(elements) * 0.5
    
    def _is_scanned_pdf(self, doc) -> bool:
        """Check if PDF appears to be a scan-only document."""
        # Sample a few pages
        pages_to_check = min(3, len(doc))
        text_count = 0
        
        for i in range(pages_to_check):
            page = doc[i]
            text = page.get_text()
            if len(text.strip()) > 50:  # Arbitrary threshold
                text_count += 1
        
        # If most sampled pages have meaningful text, it's probably not scan-only
        return text_count < pages_to_check / 2
    
    def _process_with_ocr(self, file_path: str) -> List[DocumentElement]:
        """Process a PDF with OCR using pytesseract."""
        if not HAS_OCR:
            self.logger.error("OCR processing requested but pytesseract is not installed")
            return []
        
        elements = []
        position = 0
        
        try:
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                # Convert page to image
                pix = page.get_pixmap(dpi=300)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                
                # Apply OCR
                ocr_text = pytesseract.image_to_string(img)
                
                # Process extracted text
                if ocr_text.strip():
                    # Split into paragraphs
                    paragraphs = re.split(r'\n\s*\n', ocr_text)
                    for para in paragraphs:
                        para = para.strip()
                        if not para:
                            continue
                        
                        # Detect element type (list, heading, etc.)
                        element_type = 'paragraph'  # Default
                        
                        # Check for heading
                        if len(para) < 100 and para.isupper():
                            element_type = 'heading'
                        # Check for list
                        elif self._is_list_item(para):
                            element_type = 'list'
                            
                        element = DocumentElement(
                            content=para,
                            element_type=element_type,
                            page_num=page_num + 1,
                            position=position,
                            metadata={}
                        )
                        
                        # Add list items to metadata if it's a list
                        if element_type == 'list':
                            items = self._extract_list_items(para)
                            element.metadata['items'] = items
                            element.metadata['list_type'] = 'bullet' if para.lstrip()[0] in '•*-' else 'numbered'
                        
                        elements.append(element)
                        position += 1
            
            doc.close()
            
        except Exception as e:
            self.logger.error(f"Error in OCR processing: {str(e)}", exc_info=True)
        
        return elements
    
    def _process_toc(self, toc: List) -> DocumentElement:
        """Process table of contents from PyMuPDF."""
        entries = []
        for item in toc:
            level, title, page = item[:3]
            entries.append({
                'level': level,
                'title': title,
                'page': page
            })
        
        return DocumentElement(
            content="Table of Contents",
            element_type="toc",
            page_num=0,  # TOC spans multiple pages
            position=0,
            metadata={
                'entries': entries
            }
        )
    
    def _process_blocks(self, blocks: List[Dict], page_num: int) -> List[DocumentElement]:
        """Process text blocks from PyMuPDF extraction."""
        elements = []
        position = 0
        
        # Sort blocks by y-position (top to bottom)
        blocks = sorted(blocks, key=lambda b: b["bbox"][1] if "bbox" in b else 0)
        
        for block in blocks:
            if block["type"] != 0:  # Skip non-text blocks
                continue
            
            if "lines" not in block:
                continue
            
            # Extract text from the block
            text = ""
            for line in block.get("lines", []):
                line_text = ""
                for span in line.get("spans", []):
                    line_text += span.get("text", "")
                if line_text:
                    text += line_text + " "
            
            text = text.strip()
            if not text:
                continue
            
            # Determine element type
            element_type = 'paragraph'  # Default
            metadata = {}
            
            # Store position info in metadata for header/footer detection
            metadata['bbox'] = block.get("bbox", [0, 0, 0, 0])
            
            # Get font info from first span for classification
            if block["lines"] and block["lines"][0]["spans"]:
                first_span = block["lines"][0]["spans"][0]
                font_size = first_span.get("size", 0)
                is_bold = first_span.get("font", "").lower().find("bold") >= 0
                
                # Store font information in metadata
                metadata['font_size'] = font_size
                metadata['is_bold'] = is_bold
                
                # Classify as heading
                if (is_bold and font_size > 10) or font_size > 14:
                    element_type = 'heading'
                    metadata['level'] = 1 if font_size > 16 else 2
            
            # Check for list items
            if self._is_list_item(text):
                element_type = 'list'
                items = self._extract_list_items(text)
                metadata['items'] = items
                metadata['list_type'] = 'bullet' if text.lstrip()[0] in '•*-' else 'numbered'
            
            # Check for table of contents entry
            elif self._is_toc_entry(text):
                element_type = 'toc_entry'
                parts = text.rsplit('.', 1)
                if len(parts) == 2:
                    title = parts[0].strip()
                    page = parts[1].strip()
                    metadata['title'] = title
                    metadata['page'] = page
            
            element = DocumentElement(
                content=text,
                element_type=element_type,
                page_num=page_num,
                position=position,
                metadata=metadata
            )
            elements.append(element)
            position += 1
        
        return elements
    
    def _filter_headers_footers(self, elements: List[DocumentElement], page_height: float) -> List[DocumentElement]:
        """Filter out headers and footers based on position."""
        if not elements:
            return []
        
        # Consider the top 10% of the page as header area and bottom 10% as footer area
        header_zone = page_height * 0.1
        footer_zone = page_height * 0.9
        
        # Filter elements
        filtered = []
        for element in elements:
            # We need the y-position in the original page to determine if it's a header/footer
            y_pos = element.metadata.get('bbox', [0, 0, 0, 0])[1] if 'bbox' in element.metadata else None
            
            # Skip likely headers and footers
            if y_pos is not None:
                if y_pos < header_zone and len(element.content) < 100:
                    # This is likely a header
                    element.element_type = 'header'  # Mark as header instead of removing
                elif y_pos > footer_zone and len(element.content) < 100:
                    # This is likely a footer
                    element.element_type = 'footer'  # Mark as footer instead of removing
            
            filtered.append(element)
        
        return filtered
    
    def _parse_docx_elements(self, file_path: str) -> List[DocumentElement]:
        """
        Parse DOCX document using python-docx.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            List of extracted document elements
        """
        self.logger.info(f"Parsing DOCX document: {file_path}")
        elements = []
        position = 0
        
        try:
            doc = docx.Document(file_path)
            
            # Process document structure
            # First check for TOC
            has_toc = False
            for i, para in enumerate(doc.paragraphs[:20]):  # Check first 20 paragraphs
                if para.text.lower().strip() in ["table of contents", "contents", "toc"]:
                    has_toc = True
                    toc_entries = []
                    
                    # Collect TOC entries
                    j = i + 1
                    while j < len(doc.paragraphs) and j < i + 50:  # Look at next 50 paragraphs at most
                        if self._is_toc_entry(doc.paragraphs[j].text):
                            entry_text = doc.paragraphs[j].text
                            parts = entry_text.strip().rsplit('.', 1)
                            if len(parts) == 2:
                                title = parts[0].strip()
                                page = parts[1].strip()
                                toc_entries.append({
                                    'title': title,
                                    'page': page
                                })
                            j += 1
                        else:
                            # Break if we hit a non-TOC paragraph
                            if not doc.paragraphs[j].text.strip():
                                j += 1
                                continue
                            break
                    
                    # Create TOC element
                    if toc_entries:
                        toc_element = DocumentElement(
                            content="Table of Contents",
                            element_type="toc",
                            page_num=0,
                            position=position,
                            metadata={
                                'entries': toc_entries
                            }
                        )
                        elements.append(toc_element)
                        position += 1
                        
                        # Skip these paragraphs
                        i = j
                    break
            
            # Process paragraphs
            i = 0
            while i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                text = para.text.strip()
                
                if not text:
                    i += 1
                    continue
                
                # Determine element type
                element_type = 'paragraph'  # Default
                metadata = {}
                
                # Store style information
                if para.style:
                    metadata['style'] = para.style.name
                
                # Check for heading
                if para.style and para.style.name.startswith('Heading'):
                    element_type = 'heading'
                    # Extract heading level from style name
                    level_match = re.search(r'Heading (\d+)', para.style.name)
                    level = int(level_match.group(1)) if level_match else 1
                    metadata['level'] = level
                
                # Check for list
                elif (para.style and para.style.name.startswith('List')) or self._is_list_item(text):
                    # Collect all items in this list
                    list_items = [text]
                    j = i + 1
                    
                    while j < len(doc.paragraphs):
                        next_para = doc.paragraphs[j]
                        next_text = next_para.text.strip()
                        
                        if not next_text:
                            j += 1
                            continue
                            
                        if (next_para.style and next_para.style.name.startswith('List')) or self._is_list_item(next_text):
                            list_items.append(next_text)
                            j += 1
                        else:
                            break
                    
                    element_type = 'list'
                    metadata['items'] = list_items
                    metadata['list_type'] = 'numbered' if re.match(r'^\d+\.', text) else 'bullet'
                    
                    # Create list element
                    element = DocumentElement(
                        content='\n'.join(list_items),
                        element_type=element_type,
                        page_num=0,  # DOCX doesn't have pages
                        position=position,
                        metadata=metadata
                    )
                    elements.append(element)
                    position += 1
                    
                    # Skip processed paragraphs
                    i = j
                    continue
                
                # Create element for current paragraph
                element = DocumentElement(
                    content=text,
                    element_type=element_type,
                    page_num=0,  # DOCX doesn't have pages
                    position=position,
                    metadata=metadata
                )
                elements.append(element)
                position += 1
                i += 1
            
            # Process tables
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    rows.append(row_cells)
                
                # Only add table if it has content
                if rows and any(any(cell for cell in row) for row in rows):
                    table_element = DocumentElement(
                        content=self._format_table_content(rows),  # Create readable content
                        element_type="table",
                        page_num=0,  # DOCX doesn't have pages
                        position=position,
                        metadata={
                            'rows': rows
                        }
                    )
                    elements.append(table_element)
                    position += 1
            
        except Exception as e:
            self.logger.error(f"Error in DOCX parsing: {str(e)}", exc_info=True)
        
        return self._combine_related_elements(elements)
    
    def _format_table_content(self, rows: List[List[str]]) -> str:
        """Format table rows as text for content field."""
        if not rows:
            return ""
        
        result = []
        for row in rows:
            result.append(" | ".join(cell for cell in row))
        
        return "\n".join(result)
    
    def _is_list_item(self, text: str) -> bool:
        """Check if a line of text is a list item."""
        text = text.strip()
        if not text:
            return False
            
        # Check for bullet points
        bullet_match = re.match(r'^\s*[•\-\*\+○◦➢➣➤►▶→➥➔❖]\s+', text)
        if bullet_match:
            return True
            
        # Check for numbered list patterns
        number_match = re.match(r'^\s*(\d+[\.\)]\s+|\([a-z\d]\)\s+|[a-z\d][\.\)]\s+|[ivxIVX]+[\.\)]\s+)', text)
        if number_match:
            return True
            
        return False
    
    def _extract_list_items(self, text: str) -> List[str]:
        """Extract individual list items from text."""
        items = []
        current_item = None
        
        for line in text.split('\n'):
            line = line.strip()
            if not line:
                continue
                
            if self._is_list_item(line):
                if current_item:
                    items.append(current_item)
                current_item = line
            elif current_item:
                # Continuation of previous item
                current_item += ' ' + line
            else:
                # Not part of a list
                current_item = line
        
        if current_item:
            items.append(current_item)
            
        return items
    
    def _is_toc_entry(self, text: str) -> bool:
        """Check if text is a table of contents entry."""
        text = text.strip()
        if not text:
            return False
            
        # Common TOC patterns
        has_dotted_line = '..' in text
        has_page_number = bool(re.search(r'\d+\s*$', text))
        
        return has_dotted_line and has_page_number
    
    def _clean_table(self, table: List[List[str]]) -> List[List[str]]:
        """Clean and normalize table data."""
        if not table:
            return []
            
        # Remove empty rows
        table = [row for row in table if any(cell and str(cell).strip() for cell in row)]
        
        # Normalize cells
        clean_table = []
        for row in table:
            clean_row = []
            for cell in row:
                clean_cell = str(cell).strip() if cell is not None else ""
                clean_row.append(clean_cell)
            clean_table.append(clean_row)
            
        return clean_table
    
    def _combine_related_elements(self, elements: List[DocumentElement]) -> List[DocumentElement]:
        """Combine related elements like consecutive list items."""
        if not elements:
            return []
            
        result = []
        i = 0
        
        while i < len(elements):
            current = elements[i]
            
            # Check for consecutive TOC entries
            if current.element_type == 'toc_entry':
                toc_entries = [current.metadata]
                j = i + 1
                
                while j < len(elements) and elements[j].element_type == 'toc_entry':
                    toc_entries.append(elements[j].metadata)
                    j += 1
                
                # Create combined TOC element
                if len(toc_entries) > 1:
                    toc_element = DocumentElement(
                        content="Table of Contents",
                        element_type="toc",
                        page_num=current.page_num,
                        position=current.position,
                        metadata={
                            'entries': toc_entries
                        }
                    )
                    result.append(toc_element)
                    i = j
                    continue
            
            # Check for header followed by paragraph
            if current.element_type == 'heading' and i + 1 < len(elements) and elements[i+1].element_type == 'paragraph':
                # Keep as separate elements but associate them
                next_element = elements[i+1]
                next_element.metadata['heading'] = current.content
                
                result.append(current)
                result.append(next_element)
                i += 2
                continue
            
            # Default case - keep as is
            result.append(current)
            i += 1
        
        return result
    
    def _post_process_paragraphs(self, paragraphs: List[Paragraph]) -> List[Paragraph]:
        """
        Apply additional processing to paragraphs for better integration with existing code.
        
        Args:
            paragraphs: List of paragraphs
            
        Returns:
            Processed list of paragraphs
        """
        if not paragraphs:
            return []
        
        # Combine consecutive list items if needed
        combined_lists = self._combine_lists(paragraphs)
        
        # Process headers and associate them with paragraphs
        with_headers = self._associate_headers(combined_lists)
        
        return with_headers
    
    def _combine_lists(self, paragraphs: List[Paragraph]) -> List[Paragraph]:
        """Combine consecutive list items into a single paragraph."""
        if not paragraphs:
            return []
            
        combined = []
        current_list = None
        
        for para in paragraphs:
            if para.paragraph_type == 'list':
                if current_list is None:
                    current_list = Paragraph(
                        content=para.content,
                        doc_id=para.doc_id,
                        paragraph_type='list',
                        position=para.position,
                        header_content=para.header_content
                    )
                else:
                    current_list.content += '\n' + para.content
            else:
                if current_list is not None:
                    combined.append(current_list)
                    current_list = None
                combined.append(para)
        
        # Add the last list if there is one
        if current_list is not None:
            combined.append(current_list)
        
        return combined
    
    def _associate_headers(self, paragraphs: List[Paragraph]) -> List[Paragraph]:
        """Associate headers with their following paragraphs."""
        if not paragraphs:
            return []
            
        result = []
        last_header = None
        
        for i, para in enumerate(paragraphs):
            if para.paragraph_type == 'header':
                # Add the header as a standalone paragraph
                result.append(para)
                
                # Remember this header for the next paragraph
                last_header = para.content
                
                # Associate with the next paragraph if it exists and isn't a header
                if i + 1 < len(paragraphs) and paragraphs[i + 1].paragraph_type != 'header':
                    paragraphs[i + 1].header_content = last_header
                    last_header = None  # Clear after associating
            else:
                # If this paragraph doesn't already have a header and we have a last_header
                if not para.header_content and last_header:
                    para.header_content = last_header
                    last_header = None  # Clear after associating
                
                result.append(para)
        
        return result
    
    def _is_address(self, content: str) -> bool:
        """Determine if a paragraph is an address."""
        # Check for postal code patterns
        postal_pattern = re.search(r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s?\d[A-Z]{2}\b', content)  # UK
        zip_pattern = re.search(r'\b\d{5}(-\d{4})?\b', content)  # US
        
        # Check for typical address components
        address_indicators = [
            'street', 'avenue', 'road', 'lane', 'drive', 'boulevard', 'st.', 'ave.', 'rd.',
            'apt', 'suite', 'unit', 'floor'
        ]
        has_indicator = any(indicator in content.lower() for indicator in address_indicators)
        
        # If we find postal codes or address indicators, it's likely an address
        if (postal_pattern or zip_pattern) and has_indicator:
            return True
            
        return False
    
    def _is_boilerplate(self, content: str) -> bool:
        """Determine if a paragraph is boilerplate text (footer, disclaimer, etc.)."""
        # Check for common boilerplate indicators
        boilerplate_indicators = [
            'all rights reserved', 'copyright', '©', 'confidential', 
            'disclaimer', 'terms and conditions', 'privacy policy',
            'legal notice', 'proprietary'
        ]
        
        if any(indicator in content.lower() for indicator in boilerplate_indicators):
            return True
            
        # Check for typical footer patterns (page numbers, etc.)
        if re.match(r'^\s*Page \d+ of \d+\s*$', content):
            return True
            
        return False
        
    def _setup_logger(self, level: str) -> logging.Logger:
        """Set up a logger for the document parser."""
        logger = logging.getLogger('DocumentParser')
        
        if not logger.handlers:  # Only add handlers if they don't exist
            # Set level
            level_map = {
                'DEBUG': logging.DEBUG,
                'INFO': logging.INFO,
                'WARNING': logging.WARNING,
                'ERROR': logging.ERROR
            }
            logger.setLevel(level_map.get(level.upper(), logging.INFO))
            
            # Create console handler
            console_handler = logging.StreamHandler()
            formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
            console_handler.setFormatter(formatter)
            logger.addHandler(console_handler)
        
        return logger