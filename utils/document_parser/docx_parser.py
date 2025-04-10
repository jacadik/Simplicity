"""
DOCX Parser module for extracting paragraphs from Word documents.
Enhanced with TOC handling and improved structure preservation.
"""

import time
import re
import logging
from typing import List, Dict

import docx

from .base_parser import BaseDocumentParser
from .paragraph import Paragraph

class DOCXParser(BaseDocumentParser):
    """DOCX document parser that extracts paragraphs from Word files with improved structure handling."""
    
    def parse_document(self, file_path: str, doc_id: int) -> List[Paragraph]:
        """
        Parse DOCX document and extract paragraphs.
        
        Args:
            file_path: Path to the DOCX file
            doc_id: ID of the document in the database
            
        Returns:
            List of extracted paragraphs
        """
        self.logger.info(f"Parsing DOCX document: {file_path}")
        paragraphs = []
        
        start_time = time.time()
        try:
            doc = docx.Document(file_path)
            
            # Extract paragraphs, tables and lists
            raw_paragraphs = []
            position = 0
            
            # Check for and process table of contents
            has_toc, toc_elements = self._process_toc(doc)
            if has_toc:
                # Add TOC as a single element
                raw_paragraphs.extend(toc_elements)
                position += len(toc_elements)
            
            # Process document paragraphs with improved list and header handling
            i = 0
            while i < len(doc.paragraphs):
                para = doc.paragraphs[i]
                text = para.text.strip()
                
                if not text:
                    i += 1
                    continue
                
                # Extract style information
                para_style = para.style.name if para.style else 'Normal'
                
                # Determine element type
                element_type = 'unknown'  # Default to unknown for further classification
                metadata = {'style': para_style}
                
                # Check if it's a heading
                if para_style.startswith('Heading'):
                    element_type = 'header'
                    # Extract heading level
                    level_match = re.search(r'Heading (\d+)', para_style)
                    level = int(level_match.group(1)) if level_match else 1
                    metadata['level'] = level
                
                # Check for list paragraphs and collect them as a group
                if para_style.startswith('List') or self._is_list_item(text):
                    # Collect consecutive list items
                    list_text = text
                    j = i + 1
                    
                    while j < len(doc.paragraphs):
                        next_para = doc.paragraphs[j]
                        next_text = next_para.text.strip()
                        next_style = next_para.style.name if next_para.style else 'Normal'
                        
                        if not next_text:
                            j += 1
                            continue
                            
                        if next_style.startswith('List') or self._is_list_item(next_text):
                            list_text += '\n' + next_text
                            j += 1
                        else:
                            break
                    
                    # Format the list content to ensure proper line separation
                    formatted_list_content = self._process_paragraph_content(list_text)
                    
                    element_type = 'list'
                    metadata['list_content'] = True
                    
                    # Create list element
                    raw_paragraphs.append({
                        'content': formatted_list_content,
                        'type': element_type,
                        'style': para_style,
                        'position': position,
                        'metadata': metadata
                    })
                    position += 1
                    
                    # Skip processed list items
                    i = j
                    continue
                
                # Process the paragraph content (for lists and other formatting)
                processed_content = self._process_paragraph_content(text)
                
                # Add the paragraph
                raw_paragraphs.append({
                    'content': processed_content,
                    'type': element_type,
                    'style': para_style,
                    'position': position,
                    'metadata': metadata
                })
                position += 1
                i += 1
            
            # Process tables
            table_position = position
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    rows.append(row_cells)
                
                # Only add table if it has content
                if rows and any(any(cell for cell in row) for row in rows):
                    raw_paragraphs.append({
                        'content': self._format_table(rows),
                        'type': 'table',
                        'position': table_position,
                        'metadata': {'rows': rows}
                    })
                    table_position += 1
            
            # Process extracted content with paragraph extractor
            paragraphs = self.paragraph_extractor.process_raw_paragraphs(raw_paragraphs, doc_id)
            
            elapsed_time = time.time() - start_time
            self.logger.info(f"Extracted {len(paragraphs)} paragraphs from DOCX document in {elapsed_time:.2f} seconds")
            
        except Exception as e:
            self.logger.error(f"Error parsing DOCX {file_path}: {str(e)}", exc_info=True)
        
        return paragraphs
    
    def _process_toc(self, doc) -> tuple:
        """
        Process table of contents from DOCX document.
        
        Args:
            doc: DOCX document object
            
        Returns:
            Tuple of (has_toc, toc_elements)
        """
        has_toc = False
        toc_elements = []
        
        # Check first 20 paragraphs for TOC indicators
        for i, para in enumerate(doc.paragraphs[:20]):
            if para.text.lower().strip() in ["table of contents", "contents", "toc"]:
                has_toc = True
                
                # Add TOC header
                toc_elements.append({
                    'content': para.text,
                    'type': 'toc',
                    'style': para.style.name if para.style else 'Normal',
                    'position': i,
                    'metadata': {'is_toc_header': True}
                })
                
                # Collect TOC entries
                j = i + 1
                while j < len(doc.paragraphs) and j < i + 50:  # Look at next 50 paragraphs at most
                    if self._is_toc_entry(doc.paragraphs[j].text):
                        toc_elements.append({
                            'content': doc.paragraphs[j].text,
                            'type': 'toc_entry',
                            'style': doc.paragraphs[j].style.name if doc.paragraphs[j].style else 'Normal',
                            'position': j,
                            'metadata': {'is_toc_entry': True}
                        })
                        j += 1
                    else:
                        # Break if we hit a non-TOC paragraph that's not empty
                        if doc.paragraphs[j].text.strip():
                            break
                        j += 1
                
                break
        
        return has_toc, toc_elements
    
    def _is_toc_entry(self, text: str) -> bool:
        """
        Check if text is a table of contents entry.
        
        Args:
            text: Text to check
            
        Returns:
            Boolean indicating if text is a TOC entry
        """
        text = text.strip()
        if not text:
            return False
            
        # Common TOC patterns
        has_dotted_line = '..' in text
        has_page_number = bool(re.search(r'\d+\s*$', text))
        
        return has_dotted_line and has_page_number